#!/usr/bin/env python3
"""
Pure Python Office Converter
Conversão usando bibliotecas Python nativas - SEM LibreOffice
"""

import os
import sys
from pathlib import Path
import tempfile
import io

def check_dependencies():
    """Verifica dependências necessárias"""
    missing = []
    
    try:
        import docx
        print("✅ python-docx: Disponível")
    except ImportError:
        missing.append("python-docx")
        print("❌ python-docx: Não encontrado")
    
    try:
        import openpyxl
        print("✅ openpyxl: Disponível") 
    except ImportError:
        missing.append("openpyxl")
        print("❌ openpyxl: Não encontrado")
    
    try:
        import pandas as pd
        print("✅ pandas: Disponível")
    except ImportError:
        missing.append("pandas")
        print("❌ pandas: Não encontrado")
    
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        print("✅ reportlab: Disponível")
    except ImportError:
        missing.append("reportlab")
        print("❌ reportlab: Não encontrado")
    
    return missing

def install_dependencies(packages):
    """Instala dependências automaticamente"""
    print(f"\n🔧 Instalando dependências: {', '.join(packages)}")
    
    import subprocess
    for pkg in packages:
        try:
            print(f"   Instalando {pkg}...", end=" ")
            result = subprocess.run([sys.executable, "-m", "pip", "install", pkg], 
                                  capture_output=True, text=True)
            if result.returncode == 0:
                print("✅")
            else:
                print(f"❌ {result.stderr[:50]}...")
                return False
        except Exception as e:
            print(f"❌ {e}")
            return False
    
    print("🎉 Todas as dependências instaladas!")
    return True

def docx_to_pdf_pure(docx_path, output_path):
    """Converte DOCX para PDF usando python-docx + reportlab"""
    try:
        import docx
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        
        # Ler documento DOCX
        doc = docx.Document(docx_path)
        
        # Criar PDF
        pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Adicionar título
        title_style = styles['Title']
        story.append(Paragraph("Documento Convertido", title_style))
        story.append(Spacer(1, 12))
        
        # Processar parágrafos
        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Gerar PDF
        pdf_doc.build(story)
        
        return True
        
    except Exception as e:
        print(f"❌ Erro convertendo DOCX→PDF: {e}")
        return False

def xlsx_to_csv_pure(xlsx_path, output_path):
    """Converte XLSX para CSV usando openpyxl"""
    try:
        import openpyxl
        import csv
        
        # Carregar planilha
        workbook = openpyxl.load_workbook(xlsx_path)
        worksheet = workbook.active
        
        # Escrever CSV
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            for row in worksheet.iter_rows(values_only=True):
                writer.writerow(row)
        
        return True
        
    except Exception as e:
        print(f"❌ Erro convertendo XLSX→CSV: {e}")
        return False

def csv_to_xlsx_pure(csv_path, output_path):
    """Converte CSV para XLSX usando pandas + openpyxl"""
    try:
        import pandas as pd
        
        # Ler CSV
        df = pd.read_csv(csv_path)
        
        # Salvar como XLSX
        df.to_excel(output_path, index=False, engine='openpyxl')
        
        return True
        
    except Exception as e:
        print(f"❌ Erro convertendo CSV→XLSX: {e}")
        return False

def create_test_files():
    """Cria arquivos de teste"""
    temp_dir = tempfile.mkdtemp()
    
    # 1. Criar teste DOCX
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('Teste Office Converter Pure Python', 0)
        doc.add_paragraph('Este é um teste de conversão DOCX → PDF.')
        doc.add_paragraph('Features testadas:')
        doc.add_paragraph('✅ python-docx para ler DOCX')
        doc.add_paragraph('✅ reportlab para gerar PDF')
        doc.add_paragraph('✅ Sem dependência LibreOffice')
        
        docx_path = os.path.join(temp_dir, 'teste.docx')
        doc.save(docx_path)
        print(f"✅ DOCX teste criado: {docx_path}")
    except Exception as e:
        print(f"❌ Erro criando DOCX teste: {e}")
        docx_path = None
    
    # 2. Criar teste CSV
    try:
        csv_path = os.path.join(temp_dir, 'teste.csv')
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            import csv
            writer = csv.writer(f)
            writer.writerow(['Nome', 'Idade', 'Cidade'])
            writer.writerow(['João', '30', 'São Paulo'])
            writer.writerow(['Maria', '25', 'Rio de Janeiro'])
            writer.writerow(['Pedro', '35', 'Belo Horizonte'])
        
        print(f"✅ CSV teste criado: {csv_path}")
    except Exception as e:
        print(f"❌ Erro criando CSV teste: {e}")
        csv_path = None
    
    return temp_dir, docx_path, csv_path

def benchmark_conversions(temp_dir, docx_path, csv_path):
    """Executa benchmark das conversões"""
    import time
    
    results = {}
    
    # 1. Teste DOCX → PDF
    if docx_path:
        print("\n🔄 Testando DOCX → PDF...")
        pdf_path = os.path.join(temp_dir, 'teste.pdf')
        
        start_time = time.time()
        success = docx_to_pdf_pure(docx_path, pdf_path)
        end_time = time.time()
        
        if success and os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            duration = end_time - start_time
            results['docx_to_pdf'] = {
                'success': True,
                'duration': duration,
                'size': file_size,
                'path': pdf_path
            }
            print(f"✅ Conversão DOCX→PDF: {duration:.2f}s ({file_size} bytes)")
        else:
            results['docx_to_pdf'] = {'success': False}
            print("❌ Falha DOCX→PDF")
    
    # 2. Teste CSV → XLSX
    if csv_path:
        print("\n🔄 Testando CSV → XLSX...")
        xlsx_path = os.path.join(temp_dir, 'teste.xlsx')
        
        start_time = time.time()
        success = csv_to_xlsx_pure(csv_path, xlsx_path)
        end_time = time.time()
        
        if success and os.path.exists(xlsx_path):
            file_size = os.path.getsize(xlsx_path)
            duration = end_time - start_time
            results['csv_to_xlsx'] = {
                'success': True,
                'duration': duration,
                'size': file_size,
                'path': xlsx_path
            }
            print(f"✅ Conversão CSV→XLSX: {duration:.2f}s ({file_size} bytes)")
        else:
            results['csv_to_xlsx'] = {'success': False}
            print("❌ Falha CSV→XLSX")
        
        # 3. Teste volta: XLSX → CSV
        if success:
            print("\n🔄 Testando XLSX → CSV...")
            csv2_path = os.path.join(temp_dir, 'teste_volta.csv')
            
            start_time = time.time()
            success2 = xlsx_to_csv_pure(xlsx_path, csv2_path)
            end_time = time.time()
            
            if success2 and os.path.exists(csv2_path):
                duration = end_time - start_time
                results['xlsx_to_csv'] = {
                    'success': True,
                    'duration': duration,
                    'path': csv2_path
                }
                print(f"✅ Conversão XLSX→CSV: {duration:.2f}s")
            else:
                results['xlsx_to_csv'] = {'success': False}
                print("❌ Falha XLSX→CSV")
    
    return results

def main():
    print("🐍 PURE PYTHON OFFICE CONVERTER")
    print("=" * 50)
    print("Conversão sem LibreOffice - só Python!")
    
    # 1. Verificar dependências
    print("\n1️⃣ Verificando dependências...")
    missing = check_dependencies()
    
    if missing:
        print(f"\n📦 Dependências faltando: {', '.join(missing)}")
        if input("Instalar automaticamente? (y/n): ").lower() == 'y':
            if not install_dependencies(missing):
                print("❌ Falha na instalação das dependências!")
                return
        else:
            print("💡 Instale manualmente: pip install " + " ".join(missing))
            return
    
    # 2. Criar arquivos de teste
    print("\n2️⃣ Criando arquivos de teste...")
    temp_dir, docx_path, csv_path = create_test_files()
    
    # 3. Executar benchmark
    print("\n3️⃣ Executando conversões...")
    results = benchmark_conversions(temp_dir, docx_path, csv_path)
    
    # 4. Relatório final
    print("\n" + "="*50)
    print("📊 RELATÓRIO FINAL")
    print("="*50)
    
    successful = 0
    total_time = 0
    
    for conversion, result in results.items():
        if result['success']:
            successful += 1
            duration = result.get('duration', 0)
            total_time += duration
            size = result.get('size', 0)
            print(f"✅ {conversion}: {duration:.2f}s ({size} bytes)")
        else:
            print(f"❌ {conversion}: Falhou")
    
    print(f"\n🎯 Taxa de sucesso: {successful}/{len(results)} ({successful/len(results)*100:.1f}%)")
    
    if successful > 0:
        avg_time = total_time / successful
        print(f"⏱️ Tempo médio: {avg_time:.2f}s")
        
        if avg_time < 2:
            print("🚀 Performance EXCELENTE! (Pure Python é rápido!)")
        elif avg_time < 5:
            print("✅ Performance BOA!")
        else:
            print("⚠️ Performance aceitável")
        
        print("\n🎉 PURE PYTHON OFFICE CONVERTER: FUNCIONANDO!")
        print("💡 Pronto para integração no sistema principal!")
    
    # Cleanup opcional
    import shutil
    try:
        shutil.rmtree(temp_dir)
        print(f"\n🧹 Diretório temporário removido: {temp_dir}")
    except:
        print(f"\n📁 Arquivos de teste em: {temp_dir}")

if __name__ == "__main__":
    main()